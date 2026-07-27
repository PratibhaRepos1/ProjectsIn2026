import ipaddress
import os
import json
import socket
import uuid
from typing import List
from urllib.parse import urlparse
import httpx
from sqlalchemy.orm import Session
from fastapi import UploadFile, HTTPException
from ..models.document import Document, DocumentChunk
from ..core.config import settings
from ..rag.embeddings import embed_texts


ALLOWED_TYPES = {"pdf", "docx", "txt", "csv", "xlsx", "url"}
MAX_URL_FETCH_BYTES = 5 * 1024 * 1024


def _extract_text(path: str, file_type: str) -> str:
    if file_type == "txt" or file_type == "csv":
        with open(path, encoding="utf-8", errors="ignore") as f:
            return f.read()
    if file_type == "pdf":
        try:
            import PyPDF2
            reader = PyPDF2.PdfReader(path)
            return "\n".join(p.extract_text() or "" for p in reader.pages)
        except Exception:
            return ""
    if file_type == "docx":
        try:
            import docx
            doc = docx.Document(path)
            return "\n".join(p.text for p in doc.paragraphs)
        except Exception:
            return ""
    if file_type == "xlsx":
        try:
            import openpyxl
            wb = openpyxl.load_workbook(path, data_only=True)
            lines = []
            for sheet in wb.worksheets:
                for row in sheet.iter_rows(values_only=True):
                    cells = [str(c) for c in row if c is not None]
                    if cells:
                        lines.append(" | ".join(cells))
            return "\n".join(lines)
        except Exception:
            return ""
    return ""


def _chunk_text(text: str, size: int, overlap: int) -> List[str]:
    # Chunk by line, not by a flat word list — joining an entire document's
    # words with single spaces destroys row/line structure, which silently
    # corrupts tabular data (CSV/XLSX rows bleed into each other).
    lines = text.split("\n")
    chunks: List[str] = []
    current: List[str] = []
    count = 0
    for line in lines:
        line_words = len(line.split())
        if count + line_words > size and current:
            chunks.append("\n".join(current))
            tail: List[str] = []
            tail_count = 0
            for prev_line in reversed(current):
                tail_count += len(prev_line.split())
                tail.insert(0, prev_line)
                if tail_count >= overlap:
                    break
            current, count = tail, tail_count
        current.append(line)
        count += line_words
    if current:
        chunks.append("\n".join(current))
    return chunks


def _assert_public_url(url: str) -> None:
    parsed = urlparse(url)
    if parsed.scheme not in ("http", "https"):
        raise HTTPException(status_code=400, detail="Only http/https URLs are supported")
    if not parsed.hostname:
        raise HTTPException(status_code=400, detail="Invalid URL")

    # Basic SSRF guard: a business owner could otherwise point this at
    # localhost/internal services/cloud metadata endpoints. Resolve the
    # hostname and reject anything that isn't a genuine public address.
    try:
        resolved_ip = socket.gethostbyname(parsed.hostname)
    except socket.gaierror:
        raise HTTPException(status_code=400, detail="Could not resolve URL host")

    ip = ipaddress.ip_address(resolved_ip)
    if ip.is_private or ip.is_loopback or ip.is_link_local or ip.is_reserved or ip.is_multicast:
        raise HTTPException(status_code=400, detail="URLs pointing to private/internal addresses are not allowed")


async def _fetch_url_text(url: str) -> str:
    from bs4 import BeautifulSoup

    _assert_public_url(url)
    async with httpx.AsyncClient(timeout=15, follow_redirects=True) as client:
        try:
            resp = await client.get(url, headers={"User-Agent": "ChatBizBot/1.0"})
        except httpx.HTTPError:
            raise HTTPException(status_code=400, detail="Could not fetch that URL")
    if resp.status_code >= 400:
        raise HTTPException(status_code=400, detail=f"URL returned status {resp.status_code}")
    if len(resp.content) > MAX_URL_FETCH_BYTES:
        raise HTTPException(status_code=400, detail="Page is too large")

    soup = BeautifulSoup(resp.text, "html.parser")
    for tag in soup(["script", "style", "nav", "footer", "header", "noscript"]):
        tag.decompose()
    lines = [line.strip() for line in soup.get_text(separator="\n").splitlines() if line.strip()]
    return "\n".join(lines)


async def ingest_url(db: Session, business_id: str, user_id: str, url: str) -> Document:
    text = await _fetch_url_text(url)

    doc = Document(
        business_id=business_id,
        filename=url,
        file_url=url,
        file_type="url",
        status="processing",
        uploaded_by=user_id,
    )
    db.add(doc)
    db.commit()
    db.refresh(doc)

    _embed_and_store(db, doc, text)
    return doc


async def ingest_document(
    db: Session, business_id: str, user_id: str, file: UploadFile
) -> Document:
    ext = (file.filename or "").rsplit(".", 1)[-1].lower()
    if ext not in ALLOWED_TYPES:
        raise HTTPException(status_code=400, detail=f"File type .{ext} not supported")

    content = await file.read()
    if len(content) > settings.max_upload_size_mb * 1024 * 1024:
        raise HTTPException(status_code=400, detail="File too large")

    os.makedirs(settings.upload_dir, exist_ok=True)
    saved_name = f"{uuid.uuid4()}.{ext}"
    saved_path = os.path.join(settings.upload_dir, saved_name)
    with open(saved_path, "wb") as f:
        f.write(content)

    doc = Document(
        business_id=business_id,
        filename=file.filename,
        file_url=saved_path,
        file_type=ext,
        status="processing",
        uploaded_by=user_id,
    )
    db.add(doc)
    db.commit()
    db.refresh(doc)

    _process_document(db, doc, saved_path, ext)
    return doc


def _process_document(db: Session, doc: Document, path: str, ext: str):
    text = _extract_text(path, ext)
    _embed_and_store(db, doc, text)


def _embed_and_store(db: Session, doc: Document, text: str):
    try:
        chunks = _chunk_text(text, settings.chunk_size, settings.chunk_overlap)
        embeddings = embed_texts(chunks)

        for idx, (chunk_text, emb) in enumerate(zip(chunks, embeddings)):
            chunk = DocumentChunk(
                business_id=doc.business_id,
                document_id=doc.id,
                chunk_index=idx,
                content=chunk_text,
                embedding_json=json.dumps(emb),
            )
            db.add(chunk)

        doc.status = "embedded"
    except Exception:
        doc.status = "failed"
    finally:
        db.commit()
